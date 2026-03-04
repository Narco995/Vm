import os
import re
import tempfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from config import TEMP_DIR
from bot.services.ai_service import AIService


class DocumentService:
    def __init__(self):
        self.ai = AIService()

    async def generate_document(self, topic: str, user_id: int) -> str:
        content = await self.ai.generate_report(topic)
        doc = Document()

        # Title styling
        self._set_document_margins(doc)
        title = doc.add_heading(topic, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        # Subtitle
        sub = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        doc.add_paragraph()

        # Parse and add content
        self._add_content(doc, content)

        # Save file
        filename = f"report_{user_id}_{int(datetime.now().timestamp())}.docx"
        filepath = os.path.join(TEMP_DIR, filename)
        doc.save(filepath)
        return filepath

    def _set_document_margins(self, doc):
        from docx.shared import Inches
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    def _add_content(self, doc, content: str):
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith('## ') or line.startswith('**') and line.endswith('**'):
                heading_text = line.lstrip('#').strip().strip('*')
                h = doc.add_heading(heading_text, level=2)
                h.runs[0].font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
            elif line.startswith('# '):
                heading_text = line.lstrip('#').strip()
                h = doc.add_heading(heading_text, level=1)
                h.runs[0].font.color.rgb = RGBColor(0x16, 0x3A, 0x5C)
            elif line.startswith('- ') or line.startswith('• ') or line.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                clean = line.lstrip('-•* ').strip()
                self._add_formatted_run(p, clean)
            elif re.match(r'^\d+\.', line):
                p = doc.add_paragraph(style='List Number')
                clean = re.sub(r'^\d+\.\s*', '', line)
                self._add_formatted_run(p, clean)
            else:
                p = doc.add_paragraph()
                self._add_formatted_run(p, line)

    def _add_formatted_run(self, paragraph, text: str):
        # Handle **bold** text
        parts = re.split(r'\*\*(.*?)\*\*', text)
        for i, part in enumerate(parts):
            run = paragraph.add_run(part)
            if i % 2 == 1:
                run.bold = True
            run.font.size = Pt(11)
